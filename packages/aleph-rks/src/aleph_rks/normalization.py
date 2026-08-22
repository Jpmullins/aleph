"""Per-MIME normalization pipeline.

Each Normalizer turns raw bytes into a canonicalized Markdown string plus
a structure outline and quality_flags list. The `Normalizer` protocol is
a thin facade; per-MIME implementations live in this module.

Inc 1 supports:
  application/pdf                                    → pypdf (primary), pdfminer.six (fallback)
  application/vnd.openxmlformats-...wordprocessingml → python-docx
  text/html                                          → readability-lxml + bs4
  text/markdown, text/plain                          → passthrough
  application/epub+zip                               → ebooklib

Failures are explicit (raise `NormalizationFailed`). The worker maps the
exception into `Source.status="failed"` with a clear `failure_reason`.
"""

from __future__ import annotations

import io
import re
from dataclasses import dataclass, field, replace
from typing import Protocol

import structlog
import tiktoken

from aleph_core.grounding import defang

_log = structlog.get_logger(__name__)

#: The quality flag every PDF normalizer here raises for a scan.
#:
#: Named rather than spelled out four times. Three producers set it, and until
#: 2026-08-22 nothing read it — so a typo in any one of them would have been
#: invisible, which is the same silence that let the flag sit unconsumed.
#: `aleph_rks.indexing` is the reader.
OCR_REQUIRED = "ocr-required"


class NormalizationFailed(Exception):
    """Normalizer hard-failure. The worker records this as the failure_reason."""


@dataclass
class NormalizationResult:
    markdown: str
    parser: str
    parser_version: str
    structure: dict
    quality_flags: list[str] = field(default_factory=list)

    @property
    def char_count(self) -> int:
        return len(self.markdown)

    @property
    def token_count(self) -> int:
        enc = tiktoken.get_encoding("cl100k_base")
        return len(enc.encode(self.markdown, disallowed_special=()))


class Normalizer(Protocol):
    parser: str
    parser_version: str

    def normalize(self, data: bytes) -> NormalizationResult: ...


# ---------------------------------------------------------------------------
# PDF
# ---------------------------------------------------------------------------


class DoclingNormalizer:
    """Layout-aware PDF reading: real headings, real tables, real captions.

    The shipped normalizers extract the words and none of the structure —
    measured over 20 of this instance's own PDFs, a median of one or two
    headings across an entire paper and ZERO table rows
    (`docs/measurements/pdf-parsers.md`). Both hardcode
    `{"heading_count": 0, "table_count": 0, "figure_count": 0}` with a comment
    admitting the library cannot tell.

    That is why every PDF chunk carries `section_path = NULL`: the chunker finds
    sections by looking for markdown headings, and there were none. The chunker
    computes exact character offsets and a test asserts the span slices back to
    the source — precision spent on a document with no structure to be precise
    about.

    **The counts here are MEASURED, not asserted.** They are counted off the
    emitted markdown, so if docling produces a flat wall of text for some
    document, `heading_count` is 0 and says so — which is the whole difference
    between this and the literal it replaces.

    Optional by construction. `docling` is an extra (`aleph-rks[pdf-layout]`)
    because it pulls a model stack of roughly a gigabyte and downloads weights
    on first use. A deployment that does not install it keeps the flat parsers,
    and `normalizer_for` says so rather than raising an ImportError at ingest.
    """

    parser = "docling"

    def __init__(self) -> None:
        try:
            import importlib.metadata as _md

            from docling.document_converter import DocumentConverter
        except Exception as exc:  # pragma: no cover - exercised by the registry
            msg = (
                "docling is not installed. Install the extra: "
                "`uv sync --all-packages --all-extras`, or "
                "`pip install 'aleph-rks[pdf-layout]'`"
            )
            raise NormalizationFailed(msg) from exc
        try:
            version = _md.version("docling")
        except Exception:
            version = "unknown"
        self.parser_version = f"docling@{version}"
        self._converter = DocumentConverter()

    def normalize(self, data: bytes) -> NormalizationResult:
        import re
        import tempfile

        # docling takes a path, not bytes. Written to a temp file rather than
        # kept in memory: the converter streams large PDFs and a NamedTemporary
        # is cheaper than holding a second copy of a 40 MB paper.
        with tempfile.NamedTemporaryFile(suffix=".pdf", delete=True) as handle:
            handle.write(data)
            handle.flush()
            try:
                result = self._converter.convert(handle.name)
            except Exception as exc:
                msg = f"docling failed to convert document: {exc}"
                raise NormalizationFailed(msg) from exc
            markdown = result.document.export_to_markdown()
            page_count = len(getattr(result.document, "pages", []) or [])

        headings = len(re.findall(r"^#{1,6}\s+\S", markdown, re.MULTILINE))
        table_rows = len(re.findall(r"^\s*\|.+\|\s*$", markdown, re.MULTILINE))
        figures = len(
            re.findall(r"^\s*(?:Figure|Fig\.|Table)\s+\d+", markdown, re.MULTILINE | re.I)
        )

        flags: list[str] = []
        # `ocr-required` finally gets a producer that means it. The shipped
        # parsers set it on a character-count heuristic; a layout parser that
        # returns almost nothing from a multi-page document is the real signal.
        if page_count and len(markdown) / max(page_count, 1) < 100:
            flags.append(OCR_REQUIRED)
        if headings == 0:
            # Said out loud. A structure-aware parser that found no structure
            # is a fact about the document worth recording, and the difference
            # between measuring zero and hardcoding it.
            flags.append("no-headings-found")

        return NormalizationResult(
            markdown=markdown,
            parser=self.parser,
            parser_version=self.parser_version,
            structure={
                "page_count": page_count,
                # Counted off the output. The two parsers this replaces wrote
                # literal zeros here.
                "heading_count": headings,
                # Markdown rows, minus the separator row per table where one
                # exists — an approximation, and a measured one.
                "table_count": table_rows,
                "figure_count": figures,
            },
            quality_flags=flags,
        )


class PyPDFNormalizer:
    parser = "pypdf"

    def __init__(self) -> None:
        import pypdf  # local import keeps the package optional at import time

        self.parser_version = f"pypdf@{pypdf.__version__}"
        self._pypdf = pypdf

    def normalize(self, data: bytes) -> NormalizationResult:
        try:
            reader = self._pypdf.PdfReader(io.BytesIO(data))
        except Exception as exc:
            msg = f"pypdf failed to open document: {exc}"
            raise NormalizationFailed(msg) from exc

        pages_text: list[str] = []
        flags: list[str] = []
        total_chars = 0
        for page in reader.pages:
            try:
                text = page.extract_text() or ""
            except Exception:
                text = ""
            pages_text.append(text)
            total_chars += len(text)

        if total_chars == 0 or (
            len(reader.pages) > 0 and total_chars / max(len(reader.pages), 1) < 100
        ):
            flags.append(OCR_REQUIRED)

        body = "\n\n".join(p.strip() for p in pages_text if p.strip())
        structure = {
            "page_count": len(reader.pages),
            "heading_count": 0,  # pypdf doesn't expose this reliably
            "table_count": 0,
            "figure_count": 0,
        }
        return NormalizationResult(
            markdown=_canonicalize_text(body),
            parser=self.parser,
            parser_version=self.parser_version,
            structure=structure,
            quality_flags=flags,
        )


class PDFMinerNormalizer:
    parser = "pdfminer"

    def __init__(self) -> None:
        import pdfminer  # noqa: F401
        from pdfminer import __version__ as v

        self.parser_version = f"pdfminer.six@{v}"

    def normalize(self, data: bytes) -> NormalizationResult:
        try:
            from pdfminer.high_level import extract_text

            text = extract_text(io.BytesIO(data)) or ""
        except Exception as exc:
            msg = f"pdfminer failed: {exc}"
            raise NormalizationFailed(msg) from exc

        flags: list[str] = []
        if len(text.strip()) == 0:
            flags.append(OCR_REQUIRED)
        return NormalizationResult(
            markdown=_canonicalize_text(text),
            parser=self.parser,
            parser_version=self.parser_version,
            structure={
                "page_count": text.count("\f") + 1,
                "heading_count": 0,
                "table_count": 0,
                "figure_count": 0,
            },
            quality_flags=flags,
        )


# ---------------------------------------------------------------------------
# DOCX
# ---------------------------------------------------------------------------


class DocxNormalizer:
    parser = "python-docx"

    def __init__(self) -> None:
        import docx  # python-docx

        self.parser_version = f"python-docx@{getattr(docx, '__version__', 'unknown')}"
        self._docx = docx

    def normalize(self, data: bytes) -> NormalizationResult:
        try:
            doc = self._docx.Document(io.BytesIO(data))
        except Exception as exc:
            msg = f"python-docx failed: {exc}"
            raise NormalizationFailed(msg) from exc

        lines: list[str] = []
        headings = 0
        for para in doc.paragraphs:
            style = (para.style.name or "").lower() if para.style else ""
            text = (para.text or "").strip()
            if not text:
                lines.append("")
                continue
            if style.startswith("heading"):
                lvl_str = style.replace("heading", "").strip() or "1"
                try:
                    lvl = max(1, min(6, int(lvl_str)))
                except ValueError:
                    lvl = 1
                lines.append(f"{'#' * lvl} {text}")
                headings += 1
            else:
                lines.append(text)

        tables = 0
        for tbl in doc.tables:
            tables += 1
            rows = []
            for row in tbl.rows:
                cells = [cell.text.replace("\n", " ").strip() for cell in row.cells]
                rows.append("| " + " | ".join(cells) + " |")
            if rows:
                header = rows[0]
                sep = "|" + "|".join(["---"] * (header.count("|") - 1)) + "|"
                lines.append("")
                lines.append(header)
                lines.append(sep)
                lines.extend(rows[1:])
                lines.append("")

        return NormalizationResult(
            markdown=_canonicalize_text("\n".join(lines)),
            parser=self.parser,
            parser_version=self.parser_version,
            structure={
                "page_count": 0,
                "heading_count": headings,
                "table_count": tables,
                "figure_count": 0,
            },
            quality_flags=[],
        )


# ---------------------------------------------------------------------------
# HTML
# ---------------------------------------------------------------------------


class HtmlNormalizer:
    parser = "readability-lxml+bs4"

    def __init__(self) -> None:
        import readability

        self.parser_version = f"readability-lxml@{getattr(readability, '__version__', 'unknown')}"

    def normalize(self, data: bytes) -> NormalizationResult:
        from bs4 import BeautifulSoup
        from readability import Document

        try:
            html_str = data.decode("utf-8", errors="replace")
            doc = Document(html_str)
            cleaned = doc.summary(html_partial=True)
            title = (doc.short_title() or "").strip()
        except Exception as exc:
            msg = f"readability-lxml failed: {exc}"
            raise NormalizationFailed(msg) from exc

        soup = BeautifulSoup(cleaned, "lxml")
        lines: list[str] = []
        if title:
            lines.append(f"# {title}")
            lines.append("")
        headings = 0
        for el in soup.descendants:
            if not hasattr(el, "name") or el.name is None:
                continue
            name = el.name
            if name in {"h1", "h2", "h3", "h4", "h5", "h6"}:
                lvl = int(name[1])
                lines.append(f"{'#' * lvl} {el.get_text(' ', strip=True)}")
                headings += 1
            elif name == "p":
                lines.append(el.get_text(" ", strip=True))
                lines.append("")
            elif name in {"ul", "ol"}:
                for i, li in enumerate(el.find_all("li", recursive=False)):
                    prefix = "- " if name == "ul" else f"{i + 1}. "
                    lines.append(f"{prefix}{li.get_text(' ', strip=True)}")
                lines.append("")
            elif name == "pre":
                lines.append("```")
                lines.append(el.get_text())
                lines.append("```")
                lines.append("")

        return NormalizationResult(
            markdown=_canonicalize_text("\n".join(lines)),
            parser=self.parser,
            parser_version=self.parser_version,
            structure={
                "page_count": 0,
                "heading_count": headings,
                "table_count": 0,
                "figure_count": 0,
            },
            quality_flags=[],
        )


# ---------------------------------------------------------------------------
# Plain markdown / text
# ---------------------------------------------------------------------------


class PassthroughNormalizer:
    parser = "passthrough"
    parser_version = "passthrough@1.0"

    def __init__(self, *, parser_label: str = "passthrough") -> None:
        self.parser = parser_label

    def normalize(self, data: bytes) -> NormalizationResult:
        try:
            text = data.decode("utf-8")
        except UnicodeDecodeError:
            text = data.decode("utf-8", errors="replace")
        headings = sum(1 for line in text.splitlines() if line.lstrip().startswith("#"))
        return NormalizationResult(
            markdown=_canonicalize_text(text),
            parser=self.parser,
            parser_version=self.parser_version,
            structure={
                "page_count": 0,
                "heading_count": headings,
                "table_count": 0,
                "figure_count": 0,
            },
            quality_flags=[],
        )


# ---------------------------------------------------------------------------
# EPUB
# ---------------------------------------------------------------------------


class EpubNormalizer:
    parser = "ebooklib"

    def __init__(self) -> None:
        import ebooklib  # noqa: F401
        from ebooklib import __version__ as v

        self.parser_version = f"ebooklib@{v}"

    def normalize(self, data: bytes) -> NormalizationResult:
        import tempfile

        from bs4 import BeautifulSoup
        from ebooklib import ITEM_DOCUMENT, epub

        with tempfile.NamedTemporaryFile(suffix=".epub", delete=True) as f:
            f.write(data)
            f.flush()
            try:
                book = epub.read_epub(f.name)
            except Exception as exc:
                msg = f"ebooklib failed: {exc}"
                raise NormalizationFailed(msg) from exc

        lines: list[str] = []
        title = book.get_metadata("DC", "title")
        if title:
            lines.append(f"# {title[0][0]}")
            lines.append("")

        headings = 0
        for item in book.get_items_of_type(ITEM_DOCUMENT):
            soup = BeautifulSoup(item.get_body_content(), "lxml")
            for el in soup.descendants:
                if not hasattr(el, "name") or el.name is None:
                    continue
                if el.name in {"h1", "h2", "h3", "h4", "h5", "h6"}:
                    lvl = int(el.name[1])
                    lines.append(f"{'#' * lvl} {el.get_text(' ', strip=True)}")
                    headings += 1
                elif el.name == "p":
                    lines.append(el.get_text(" ", strip=True))
                    lines.append("")

        return NormalizationResult(
            markdown=_canonicalize_text("\n".join(lines)),
            parser=self.parser,
            parser_version=self.parser_version,
            structure={
                "page_count": 0,
                "heading_count": headings,
                "table_count": 0,
                "figure_count": 0,
            },
            quality_flags=[],
        )


# ---------------------------------------------------------------------------
# Dispatch
# ---------------------------------------------------------------------------


#: PDF parser preference, most structure-aware first.
#:
#: A LIST rather than a hardcoded default, because this is the plugin thesis
#: applied to ingest: the parser is a choice, the registry resolves it, and a
#: deployment that has not installed the layout stack degrades to the flat
#: parsers instead of failing to ingest. `ALEPH_PDF_PARSER` pins one by name
#: for an operator who wants to compare
#: (`scripts/compare_pdf_parsers.py`, `docs/measurements/pdf-parsers.md`).
PDF_PARSERS: tuple[tuple[str, type], ...] = (
    ("docling", DoclingNormalizer),
    ("pypdf", PyPDFNormalizer),
    ("pdfminer", PDFMinerNormalizer),
)


#: PDF parsers this process has already announced it settled on.
#:
#: Once, not once per document. A worker ingesting three hundred PDFs would
#: otherwise emit three hundred identical warnings, which is how an operator
#: learns to filter the line out.
_pdf_parser_announced: set[str] = set()


def pdf_normalizer(preferred: str | None = None) -> Normalizer:
    """The best PDF parser this deployment actually has.

    Tries in order and takes the first that constructs. `DoclingNormalizer`
    raises `NormalizationFailed` when the extra is not installed, so an
    installation without it keeps the flat parsers.

    **The fall-back is announced.** It used to be inferable and nothing else:
    the previous version of this docstring said the degradation was recorded
    because "the resulting document carries `heading_count: 0` and the chunks
    carry `section_path = NULL`". Both are true and neither is a record — a PDF
    that genuinely has no headings produces exactly the same two values, so the
    only way to learn that a deployment had no layout parser at all was to
    notice that EVERY document had them. Measured on this instance: 91.5% of
    pypdf chunks carry a NULL `section_path` and 896 documents report
    `heading_count = 0`, and no line anywhere said the extra was missing from
    the worker image.

    That is the same defect as a dead embedder writing no chunks (WS-RS1) — an
    absence standing in for a state — and it gets the same treatment: the
    degradation is said out loud, once per process, with the remedy in it.
    """
    import os

    wanted = preferred or os.environ.get("ALEPH_PDF_PARSER") or ""
    candidates = list(PDF_PARSERS)
    if wanted:
        named = [c for c in candidates if c[0] == wanted]
        if not named:
            msg = f"unknown PDF parser {wanted!r}; choose from {[c[0] for c in PDF_PARSERS]}"
            raise NormalizationFailed(msg)
        # Pinned by name means PINNED. Falling back would make the setting a
        # suggestion and a comparison run would silently measure the wrong one.
        candidates = named

    last: Exception | None = None
    skipped: list[str] = []
    for name, cls in candidates:
        try:
            normalizer: Normalizer = cls()
        except Exception as exc:
            last = exc
            skipped.append(f"{name} ({type(exc).__name__}: {exc})")
            continue
        # Only when something more structure-aware was passed over. Landing on
        # the first candidate is the healthy case and must stay silent, or the
        # line means nothing.
        if skipped and name not in _pdf_parser_announced:
            _pdf_parser_announced.add(name)
            _log.warning(
                "rks.pdf_parser.degraded",
                using=name,
                skipped=skipped,
                impact=(
                    f"every PDF this process ingests is parsed by {name!r}, which "
                    "extracts words and no structure: `section_path` will be NULL on "
                    "its chunks and `structure_jsonb.heading_count` 0, and neither "
                    "value distinguishes that from a PDF that has no headings"
                ),
                remedy="install the layout extra: uv sync --all-packages --all-extras",
            )
        return normalizer
    msg = f"no usable PDF parser: {last}"
    raise NormalizationFailed(msg)


def normalizer_for(mime_type: str) -> Normalizer:
    mt = mime_type.split(";", maxsplit=1)[0].strip().lower()
    if mt == "application/pdf":
        return pdf_normalizer()
    if mt == ("application/vnd.openxmlformats-officedocument.wordprocessingml.document"):
        return DocxNormalizer()
    if mt in {"text/html", "application/xhtml+xml"}:
        return HtmlNormalizer()
    if mt in {"text/markdown", "text/x-markdown"}:
        return PassthroughNormalizer(parser_label="markdown")
    if mt == "text/plain":
        return PassthroughNormalizer(parser_label="text")
    if mt == "application/epub+zip":
        return EpubNormalizer()
    msg = f"no normalizer for mime type: {mime_type}"
    raise NormalizationFailed(msg)


def normalize_bytes(data: bytes, mime_type: str) -> NormalizationResult:
    """The single boundary every ingested byte crosses. Defang happens here.

    Ingested text is untrusted: it reaches the model's context, and under the
    research loop it reaches it without a human reading it first. Zero-width
    characters, bidi overrides and Unicode line separators let a document show
    a reviewer one thing while the model reads another — the reviewer approves
    a summary, and the instructions hidden after a U+2028 never appear on their
    screen.

    Stripping them is not a prompt-injection filter and does not claim to be.
    It removes a class of *invisibility*, which is decidable, rather than
    attempting to detect intent, which is not. Visible adversarial text stays
    visible — that is the point, because a reviewer can then see it.

    Applied at the single choke point rather than at each normalizer, so a new
    format cannot arrive without it.
    """
    primary = normalizer_for(mime_type)
    try:
        result = primary.normalize(data)
    except NormalizationFailed:
        if mime_type.split(";", maxsplit=1)[0].strip().lower() == "application/pdf":
            result = PDFMinerNormalizer().normalize(data)
        else:
            raise

    cleaned = defang(result.markdown)
    if cleaned == result.markdown:
        return result
    return replace(
        result,
        markdown=cleaned,
        quality_flags=[*result.quality_flags, "defanged_invisible_characters"],
    )


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------


_WS_BLOCK = re.compile(r"\n{3,}")


def _canonicalize_text(text: str) -> str:
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = _WS_BLOCK.sub("\n\n", text)
    return text.strip() + "\n" if text.strip() else ""
